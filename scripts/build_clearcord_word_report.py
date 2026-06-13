from __future__ import annotations

import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_SCRIPTS_DIR = Path(
    r"C:\Users\Admin\.codex\plugins\cache\openai-primary-runtime\documents\26.601.10930\skills\documents\scripts"
)
if str(SKILL_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SKILL_SCRIPTS_DIR))

from table_geometry import apply_table_geometry  # type: ignore  # noqa: E402


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "generated"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUTPUT_DIR / "ClearCord_KienTruc_va_DacTaChucNang.docx"
DIAGRAM_PATH = OUTPUT_DIR / "clearcord_architecture.png"

REPORT_DATE = "13/06/2026"
TITLE = "ClearCord - Tai lieu kien truc va dac ta chuc nang"
SUBTITLE = "Tong hop tu source code hien tai cua he thong web chat ClearCord"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "E8EEF5"
BORDER = "C9D2DF"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_paragraph_border_bottom(paragraph, color: str = "D7DBE2", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.color.rgb = INK
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    if "Subtitle" not in styles:
        subtitle = styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = GRAY
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def add_header_footer(document: Document) -> None:
    section = document.sections[0]

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ClearCord | Kien truc he thong va dac ta chuc nang")
    set_run_font(run, size=9, color=GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"Bao cao tong hop source code | {REPORT_DATE}")
    set_run_font(run, size=9, color=GRAY)


def add_cover(document: Document) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(TITLE)
    set_run_font(run, size=24, color=INK, bold=True)

    p = document.add_paragraph(style="Subtitle")
    run = p.add_run(SUBTITLE)
    set_run_font(run, size=13, color=GRAY)

    metadata = document.add_table(rows=4, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [2700, 6660]
    apply_table_geometry(metadata, widths, table_width_dxa=9360, indent_dxa=120)

    rows = [
        ("Du an", "ClearCord"),
        ("Pham vi", "Frontend React/Vite, backend ASP.NET Core, realtime SignalR, voice/video WebRTC, Clear AI"),
        ("Nguon doi chieu", "Controllers, services, repositories, SignalR hub, frontend pages/components va cau hinh hien co trong workspace"),
        ("Ngay tong hop", REPORT_DATE),
    ]

    for idx, (label, value) in enumerate(rows):
        metadata.cell(idx, 0).text = label
        metadata.cell(idx, 1).text = value
        for col in (0, 1):
            cell = metadata.cell(idx, col)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(metadata.cell(idx, 0), HEADER_FILL)
        label_run = metadata.cell(idx, 0).paragraphs[0].runs[0]
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = metadata.cell(idx, 1).paragraphs[0].runs[0]
        set_run_font(value_run, size=10.5, color=GRAY)

    document.add_paragraph(
        "Tai lieu nay gom hai phan: so do kien truc tong the cua he thong web ClearCord va dac ta day du cac chuc nang dang co, kem giai thich cach Clear AI hoat dong va cong nghe tao nen lop AI do."
    )

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(8)
    set_paragraph_border_bottom(rule)


def add_heading(document: Document, text: str, level: int) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_body(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, size=11, color=INK, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=11, color=GRAY if bold_prefix else None)


def add_simple_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)

    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border(cell)
        run = cell.paragraphs[0].runs[0]
        set_run_font(run, size=10.5, color=INK, bold=True)

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            run = cell.paragraphs[0].runs[0]
            set_run_font(run, size=10.25, color=GRAY)


def architecture_components() -> list[list[str]]:
    return [
        ["Trinh duyet / client", "React 18, Vite, JavaScript, CSS", "Render giao dien SPA, xu ly state giao dien, goi API va nhan su kien realtime."],
        ["Lop giao tiep API", "Axios + REST API", "Xu ly CRUD cho auth, user, friends, servers, channels, messages, notifications va AI."],
        ["Lop realtime", "@microsoft/signalr + ASP.NET SignalR Hub", "Phat su kien messageCreated, typingChanged, presenceChanged, notificationCreated, voiceParticipantsUpdated va webrtcSignal."],
        ["Lop nghiep vu backend", "ASP.NET Core Services", "Ap dung logic dang nhap, phan quyen, tao server, xu ly tin nhan, thong bao, voice va Clear AI."],
        ["Lop du lieu", "Repositories + EF Core DbContext", "Truy cap SQL Server theo mo hinh entities, DTOs va repository pattern."],
        ["Co so du lieu", "SQL Server / LocalDB", "Luu users, friendships, servers, roles, channels, messages, attachments, notifications, voice participants va revoked tokens."],
        ["Lop AI", "ClearAiService + OpenAI-compatible API", "Phan tich lenh noi bo, suy dien ngu canh chat va goi model ngoai khi duoc cau hinh."],
        ["Voice / video media", "WebRTC + Browser Media APIs", "Truyen media peer-to-peer, lay microphone/camera/screen stream va dong bo state qua SignalR."],
    ]


def feature_modules() -> list[tuple[str, list[list[str]]]]:
    return [
        (
            "1. Xac thuc va phien dang nhap",
            [
                ["Dang ky", "Tao tai khoan moi va tra ve access token cung thong tin user.", "ASP.NET Identity, JWT, Axios"],
                ["Dang nhap", "Xac thuc email/ten dang nhap va mat khau, luu token vao localStorage.", "JWT Bearer, Axios, localStorage"],
                ["Dang xuat", "Thu hoi token bang co che revoked token va xoa session tren client.", "RevokedToken repository, JWT events"],
                ["Quen / dat lai mat khau", "Sinh reset token va doi mat khau thong qua endpoint reset.", "Identity token providers"],
            ],
        ),
        (
            "2. Ho so nguoi dung",
            [
                ["Xem thong tin ca nhan", "Lay profile hien tai va profile theo user id.", "UsersController, React modals"],
                ["Cap nhat profile", "Sua display name, thong tin ho so va dong bo lai client session.", "UserService, Axios"],
                ["Tai avatar", "Upload anh dai dien bang multipart form data.", "IFormFile, FileStorageService"],
                ["Tim kiem user", "Tim nguoi dung de ket ban, nhan tin hoac xem profile.", "EF Core query + search endpoint"],
            ],
        ),
        (
            "3. Ban be va direct conversation",
            [
                ["Gui loi moi ket ban", "Tao friend request den user muc tieu.", "FriendsController, SQL Server"],
                ["Chap nhan / tu choi loi moi", "Cap nhat trang thai request va dua user vao danh sach ban be.", "FriendService"],
                ["Danh sach ban be", "Hien thi online/offline, avatar va hanh dong chat/call.", "React state, presence update"],
                ["DM 1-1", "Tao hoac mo direct conversation giua hai user.", "DirectConversationService, SignalR"],
            ],
        ),
        (
            "4. Server, invite va workspace",
            [
                ["Tao server", "Tao workspace moi theo kieu Discord-style.", "ServersController, EF Core"],
                ["Tham gia bang invite", "Join server thong qua invite code.", "JoinServer endpoint"],
                ["Roi / xoa server", "Thanh vien co the roi; owner/admin du quyen xoa server.", "ServerService + permissions"],
                ["Xem chi tiet server", "Lay categories, channels, roles, members va metadata cua server.", "ServerDetailsDto"],
                ["Lay invite", "Sinh hoac lay invite link hien tai de chia se.", "Server invite endpoint"],
            ],
        ),
        (
            "5. Channel va cau truc noi dung",
            [
                ["Category", "Tao, sua, xoa category de nhom channel.", "ChannelsController, ChannelCategory entity"],
                ["Text channel", "Tao, sua, xoa kenh chat text.", "ChannelService, ChannelType.Text"],
                ["Voice channel", "Tao, sua, xoa kenh voice/video.", "ChannelService, ChannelType.Voice"],
                ["Thu tu va topic", "Quan ly position, topic, categoryId de sap xep workspace.", "DTOs + admin UI"],
            ],
        ),
        (
            "6. Tin nhan va noi dung chat",
            [
                ["Gui tin nhan kenh", "Gui message vao text channel, co the kem file va reply.", "MessagesController, multipart upload"],
                ["Gui direct message", "Gui tin nhan vao direct conversation.", "DirectConversationsController, SignalR"],
                ["Sua / xoa tin nhan", "Cap nhat noi dung hoac danh dau deleted.", "MessageService, realtime update"],
                ["Reply", "Gan quan he replyTo de hien ngu canh tra loi.", "MessageDto, UI rendering"],
                ["Reaction", "Them / go emoji reaction.", "MessageReaction entity"],
                ["Pin", "Ghim va bo ghim tin nhan quan trong.", "TogglePin endpoint"],
                ["Dinh kem tep", "Gui tep va preview anh trong khung chat.", "FileStorageService, MessageAttachment"],
            ],
        ),
        (
            "7. Realtime, presence va notifications",
            [
                ["Realtime message", "Tin nhan moi duoc day qua hub ngay sau khi tao.", "SignalR messageCreated"],
                ["Typing indicator", "Thong bao dang go trong channel hoac DM.", "typingChanged event"],
                ["Online / offline", "Theo doi ket noi va phat su kien presenceChanged.", "UserConnection, SignalR"],
                ["Thong bao he thong", "Thong bao tin nhan moi, friend request va su kien server.", "NotificationService"],
                ["Danh dau da doc", "Danh dau tung thong bao hoac tat ca la da doc.", "NotificationsController"],
            ],
        ),
        (
            "8. Voice, video va chia se man hinh",
            [
                ["Tham gia voice channel", "User vao phong voice tren server va nhan danh sach participant.", "VoiceService, SignalR"],
                ["Direct call", "Voice/video call trong direct conversation.", "DirectVoiceService, SignalR"],
                ["Mic / camera / screen state", "Dong bo mute, camera va screen sharing giua cac participant.", "UpdateVoiceState"],
                ["WebRTC signaling", "Trao doi Offer, Answer, ICE Candidate va Hangup qua hub.", "ChatHub + RTCPeerConnection"],
                ["Media peer-to-peer", "Audio/video stream chay truc tiep giua cac client, backend chi mang signaling.", "WebRTC, getUserMedia, getDisplayMedia"],
            ],
        ),
        (
            "9. Quan tri server va phan quyen",
            [
                ["Role va permission", "Tao role, to mau, cap quyen theo tap PermissionType.", "ServerRole, ServerRolePermission"],
                ["Gan / go role", "Gan role cho thanh vien va thu hoi khi can.", "AssignRole / RemoveRole"],
                ["Quan ly channel/server", "Bat buoc quyen ManageChannels hoac ManageServer.", "PermissionService"],
                ["Kick / ban member", "Moderate member bang ly do tuy chon.", "KickMembers, BanMembers permissions"],
            ],
        ),
        (
            "10. Clear AI va tro ly giong noi",
            [
                ["Wake phrase va voice shortcut", "Lang nghe wake phrase va hotkey Ctrl+Space de kich hoat.", "SpeechRecognition, browser events"],
                ["Doc tin nhan", "AI co the doc recent messages, latest message va xac dinh nguoi gui.", "Regex parser + MessageService"],
                ["Gui tin nhan bang lenh", "AI co the soan/gui DM hoac gui vao channel hien tai/channel dat ten.", "ClearAiService"],
                ["Mo voice/video call", "AI co the khoi dong intent goi voice/video voi ban be.", "ClearAiActionDto + UI action"],
                ["Fallback chat model", "Neu bat OpenAI provider, request se duoc day den chat/completions.", "HttpClient, OpenAI-compatible API"],
                ["Noi va phan hoi am thanh", "Tra loi bang text va speech synthesis ngay tren client.", "SpeechSynthesis"],
            ],
        ),
    ]


def add_architecture_section(document: Document) -> None:
    add_heading(document, "I. So do kien truc tong the", 1)
    document.add_paragraph(
        "Kien truc cua ClearCord duoc to chuc theo kieu client-server co them hai lop giao tiep thoi gian thuc va media peer-to-peer. Frontend React chay trong trinh duyet, backend ASP.NET Core cung cap REST API va SignalR Hub, SQL Server luu du lieu nghiep vu, con WebRTC dam nhiem truyen audio/video giua cac client."
    )

    document.add_picture(str(DIAGRAM_PATH), width=Inches(6.2))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("Hinh 1. So do kien truc tong the cua ClearCord")
    set_run_font(run, size=10, color=GRAY, italic=True)

    add_simple_table(
        document,
        ["Thanh phan", "Cong nghe", "Vai tro trong he thong"],
        architecture_components(),
        [2100, 2500, 4760],
    )

    add_heading(document, "II. Luong tuong tac chinh", 1)
    add_body(document, "Nguoi dung mo ung dung tu ASP.NET Core hoac Vite dev server, sau do frontend React khoi tao giao dien va phien dang nhap. ", "1. Truy cap va khoi tao: ")
    add_body(document, "Moi thao tac CRUD nhu dang nhap, cap nhat profile, tao server, tao channel, gui message kem file deu di qua REST API. ", "2. CRUD va business flows: ")
    add_body(document, "Khi can tinh thoi gian thuc, frontend ket noi SignalR Hub de nhan messageCreated, typingChanged, presenceChanged, notificationCreated va voiceParticipantsUpdated. ", "3. Realtime: ")
    add_body(document, "Voice/video media khong di qua backend; backend chi chuyen tin signaling Offer, Answer, ICE Candidate va Hangup de hai browser thiet lap ket noi WebRTC truc tiep. ", "4. Voice/video: ")
    add_body(document, "Khi nguoi dung goi Clear AI, backend uu tien parser lenh noi bo; neu lenh vuot pham vi parser va provider duoc cau hinh, he thong goi them model ngoai de sinh cau tra loi. ", "5. AI flow: ")


def add_feature_sections(document: Document) -> None:
    add_heading(document, "III. Dac ta chuc nang theo module", 1)
    document.add_paragraph(
        "Bang dac ta duoi day tong hop cac chuc nang hien dang duoc the hien trong controllers, services, hub SignalR va giao dien frontend."
    )

    for title, rows in feature_modules():
        add_heading(document, title, 2)
        add_simple_table(
            document,
            ["Chuc nang", "Mo ta nghiep vu", "Cong nghe thuc hien"],
            rows,
            [2200, 4700, 2460],
        )


def add_ai_section(document: Document) -> None:
    add_heading(document, "IV. Cach thuc AI hoat dong trong ClearCord", 1)
    add_heading(document, "1. Dau vao tu text va giong noi", 2)
    document.add_paragraph(
        "Clear AI nhan prompt tu o chat hoac tu giong noi. Tren frontend, component ClearAssistantPanel su dung Web Speech API de nghe wake phrase, nhan SpeechRecognition transcript va phat lai cau tra loi qua SpeechSynthesis."
    )

    add_heading(document, "2. Phan tich lenh noi bo", 2)
    document.add_paragraph(
        "Backend ClearAiService chuan hoa cau lenh, loai bo wake phrase, xu ly dau tieng Viet va so khop bang nhieu bieu thuc chinh quy. Dich vu nay co the nhan biet cac nhom lenh nhu doc tin nhan, doc tin nhan moi nhat, gui DM, gui vao current channel, gui vao named channel, mo compose draft va bat dau voice/video call."
    )
    document.add_paragraph(
        "De tim dung ban be hoac kenh, service con dung chuan hoa chuoi, compact lookup va Levenshtein distance de cho phep nguoi dung noi/go ten gan dung van tim duoc doi tuong."
    )

    add_heading(document, "3. Giai doan thuc thi hanh dong", 2)
    document.add_paragraph(
        "Khi parser da hieu yeu cau, Clear AI khong dung lai o muc tra loi van ban. No goi truc tiep cac service nghiep vu nhu MessageService, DirectConversationService, VoiceService va NotificationService de tao hanh dong that, sau do tra ve ClearAiActionDto de frontend mo dung man hinh hoac che do soan tin."
    )

    add_heading(document, "4. Fallback sang model ngoai", 2)
    document.add_paragraph(
        "Neu provider duoc dat la openai va co API key, service se goi endpoint chat/completions voi model cau hinh trong appsettings. Luc nay he thong gui kem prompt, language va current scope de model tra loi tu nhien hon cho nhung cau hoi khong khop parser."
    )

    add_heading(document, "5. Gioi han hien tai", 2)
    document.add_paragraph(
        "Phien ban hien tai la mot AI task-oriented trong app chat. No lam rat tot voi nhan/gui tin, doc thong tin gan day va mo call, nhung chua phai agent tong quat co kha nang suy luan ngoai pham vi nghiep vu cua ClearCord."
    )

    add_heading(document, "V. Cong nghe tao nen lop AI", 1)
    add_simple_table(
        document,
        ["Lop cong nghe", "Thanh phan cu the", "Vai tro"],
        [
            ["AI giao tiep giong noi", "SpeechRecognition, SpeechSynthesis", "Thu am lenh, nhan dang van ban va doc phan hoi tren trinh duyet."],
            ["Bo parser lenh", "Regex, Unicode normalization, string matching", "Nhan dien y dinh va tham so nghiep vu tu prompt tieng Viet/Anh."],
            ["Bo tim doi tuong gan dung", "Levenshtein distance, compact lookup", "Tim ban be, kenh va context ngay ca khi nguoi dung nhap/noi chua chinh xac tuyet doi."],
            ["Dieu phoi hanh dong", "ClearAiService, DTO actions, ASP.NET services", "Anh xa y dinh AI thanh thao tac thuc su trong he thong."],
            ["Model hoi thoai ngoai", "OpenAI-compatible chat/completions, model gpt-4.1-mini theo cau hinh", "Sinh cau tra loi tu nhien khi parser noi bo khong phu hop."],
        ],
        [2200, 3200, 3960],
    )

    add_heading(document, "VI. Ket luan ky thuat", 1)
    document.add_paragraph(
        "ClearCord la he thong web chat kieu Discord duoc xay theo mo hinh tach lop ro rang: React/Vite cho giao dien, ASP.NET Core cho API va nghiep vu, EF Core + SQL Server cho du lieu, SignalR cho realtime, WebRTC cho media va Clear AI cho dieu phoi lenh. Kien truc nay de mo rong them moderation, media server, search va AI agents o cac giai doan sau."
    )


def create_architecture_diagram(path: Path) -> None:
    width, height = 1800, 1200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    font_path = r"C:\Windows\Fonts\arial.ttf"
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    title_font = ImageFont.truetype(bold_path, 40)
    box_title_font = ImageFont.truetype(bold_path, 28)
    body_font = ImageFont.truetype(font_path, 22)
    small_font = ImageFont.truetype(font_path, 20)

    def rounded_box(x1, y1, x2, y2, fill, outline, title, lines):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline=outline, width=4)
        draw.text((x1 + 24, y1 + 18), title, fill=(11, 37, 69), font=box_title_font)
        current_y = y1 + 68
        for line in lines:
            draw.text((x1 + 28, current_y), line, fill=(70, 70, 70), font=body_font)
            current_y += 34

    def arrow(x1, y1, x2, y2, label, color=(46, 116, 181)):
        draw.line((x1, y1, x2, y2), fill=color, width=6)
        if abs(x2 - x1) > abs(y2 - y1):
            direction = 1 if x2 >= x1 else -1
            draw.polygon(
                [
                    (x2, y2),
                    (x2 - 20 * direction, y2 - 10),
                    (x2 - 20 * direction, y2 + 10),
                ],
                fill=color,
            )
        else:
            direction = 1 if y2 >= y1 else -1
            draw.polygon(
                [
                    (x2, y2),
                    (x2 - 10, y2 - 20 * direction),
                    (x2 + 10, y2 - 20 * direction),
                ],
                fill=color,
            )
        if label:
            tx = (x1 + x2) / 2
            ty = (y1 + y2) / 2 - 24
            draw.rounded_rectangle((tx - 130, ty - 18, tx + 130, ty + 18), radius=12, fill="white", outline=color, width=2)
            text_bbox = draw.textbbox((0, 0), label, font=small_font)
            tw = text_bbox[2] - text_bbox[0]
            th = text_bbox[3] - text_bbox[1]
            draw.text((tx - tw / 2, ty - th / 2), label, fill=color, font=small_font)

    draw.text((60, 40), "So do kien truc tong the ClearCord", fill=(11, 37, 69), font=title_font)
    draw.text((60, 95), "Client React/Vite + backend ASP.NET Core + realtime SignalR + media WebRTC + Clear AI", fill=(89, 89, 89), font=body_font)

    rounded_box(
        70,
        180,
        500,
        430,
        "#F7FAFD",
        "#AFC3DA",
        "Nguoi dung / Browser",
        [
            "- Dang nhap, chat, quan ly server",
            "- Nghe/nhan lenh giong noi cho Clear AI",
            "- Mic, camera, screen sharing",
        ],
    )

    rounded_box(
        610,
        150,
        1210,
        470,
        "#EEF5FB",
        "#7FA6CC",
        "Frontend SPA",
        [
            "- React 18 + Vite",
            "- Axios cho REST API",
            "- SignalR client cho realtime",
            "- VoicePanel + ClearAssistantPanel",
            "- state UI, notifications, profile, admin",
        ],
    )

    rounded_box(
        610,
        560,
        1210,
        930,
        "#F2F4F7",
        "#8A9BB0",
        "Backend ASP.NET Core",
        [
            "- Controllers: Auth, Users, Friends, Servers",
            "- Channels, Messages, Notifications, Voice, AI",
            "- ChatHub: typing, presence, WebRTC signaling",
            "- Services + Repositories + EF Core",
        ],
    )

    rounded_box(
        1320,
        580,
        1710,
        860,
        "#F9FBFD",
        "#AFC3DA",
        "SQL Server / LocalDB",
        [
            "- Users, roles, servers",
            "- Channels, messages, attachments",
            "- Notifications, voice participants",
            "- Revoked tokens / connections",
        ],
    )

    rounded_box(
        1320,
        180,
        1710,
        420,
        "#FFF8F0",
        "#E2B07A",
        "OpenAI-compatible API",
        [
            "- chat/completions",
            "- Duoc goi khi Clear AI can fallback",
            "- Model cau hinh: gpt-4.1-mini",
        ],
    )

    rounded_box(
        70,
        620,
        500,
        930,
        "#F7FBF8",
        "#9BC7AA",
        "Browser media layer",
        [
            "- Web Speech API",
            "- getUserMedia / getDisplayMedia",
            "- RTCPeerConnection",
            "- SpeechSynthesis",
        ],
    )

    rounded_box(
        1320,
        930,
        1710,
        1090,
        "#FAFAFA",
        "#B6B6B6",
        "Remote peers",
        [
            "- Audio/video stream di truc tiep",
            "- Backend chi chuyen signaling",
        ],
    )

    arrow(500, 300, 610, 300, "UX thao tac")
    arrow(910, 470, 910, 560, "REST + SignalR")
    arrow(1210, 700, 1320, 700, "EF Core")
    arrow(1210, 300, 1320, 300, "HTTP model API")
    arrow(500, 760, 610, 760, "Browser APIs")
    arrow(1040, 930, 1480, 930, "WebRTC signaling qua hub")
    arrow(500, 860, 1320, 1010, "Media P2P", color=(76, 147, 94))

    image.save(path)


def build_document() -> Path:
    create_architecture_diagram(DIAGRAM_PATH)

    document = Document()
    configure_document(document)
    add_header_footer(document)
    add_cover(document)
    add_architecture_section(document)
    add_feature_sections(document)
    add_ai_section(document)

    document.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
